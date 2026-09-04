# -*- coding: utf-8 -*-
"""文件解析工具：xlsx / docx / pdf → 结构化数据。
三步走导入流程 Step 1：仅解析预览，禁止直接写库。"""
import hashlib
import io
import re
import time
import uuid
from collections import OrderedDict
from datetime import date, datetime

import pandas as pd

from ..exceptions import BizError

# 上传文件解析结果缓存：file_id -> 解析结果（内存态，重启即失效，可重新上传）
_PARSE_CACHE: OrderedDict = OrderedDict()
_PARSE_CACHE_TTL = 30 * 60
_PARSE_CACHE_LIMIT = 20

MAX_PREVIEW_ROWS = 10  # 预览行数（文档要求前10行）


# ---------------------------------------------------------------- 文件指纹
def file_md5(data: bytes) -> str:
    return hashlib.md5(data).hexdigest()


def make_file_id() -> str:
    return uuid.uuid4().hex


def cache_parsed(file_id: str, payload: dict) -> None:
    now = time.monotonic()
    for key, (cached_at, _payload) in list(_PARSE_CACHE.items()):
        if now - cached_at > _PARSE_CACHE_TTL:
            _PARSE_CACHE.pop(key, None)
    _PARSE_CACHE[file_id] = (now, payload)
    _PARSE_CACHE.move_to_end(file_id)
    while len(_PARSE_CACHE) > _PARSE_CACHE_LIMIT:
        _PARSE_CACHE.popitem(last=False)


def get_parsed(file_id: str) -> dict:
    cached = _PARSE_CACHE.get(file_id)
    if not cached:
        raise BizError("文件解析结果已失效，请重新上传", code=404)
    cached_at, payload = cached
    if time.monotonic() - cached_at > _PARSE_CACHE_TTL:
        _PARSE_CACHE.pop(file_id, None)
        raise BizError("文件解析结果已过期，请重新上传", code=404)
    _PARSE_CACHE.move_to_end(file_id)
    return payload


# ---------------------------------------------------------------- 单元格规范化
def _cell_to_str(v):
    if v is None:
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    if isinstance(v, (datetime, date, pd.Timestamp)):
        return str(pd.Timestamp(v).date())
    return str(v).strip()


def _grid_rows(df: pd.DataFrame) -> list:
    """DataFrame(header=None 读取) → 全量字符串网格。所有值转字符串，空单元格为 ""。"""
    rows = df.astype(object).where(pd.notnull(df), "").values.tolist()
    return [[_cell_to_str(c) for c in r] for r in rows]


def _detect_header_row(grid: list) -> int:
    """自动探测列名行（0 基下标）：连续跳过"非空单元格数 < 2"的行（合并标题行只有
    左上角有值，必被跳过），取首个 ≥2 个非空单元格的行为列名行；找不到回退 0。
    首行即多格 → 结果 0，与旧 header=0 解析逐字节一致（零回归）。"""
    for i, row in enumerate(grid):
        if sum(1 for c in row if str(c).strip()) >= 2:
            return i
    return 0


def _sheet_payload(headers: list, data_rows: list) -> dict:
    """headers + 数据区 → {headers, rows, total_rows, preview_limit, _full_rows}。
    rows 仅预览前 10 行；_full_rows 缓存完整数据供确认入库阶段使用（内存态）。"""
    return {
        "headers": [str(c).strip() for c in headers],
        "rows": data_rows[:MAX_PREVIEW_ROWS],
        "total_rows": len(data_rows),
        "preview_limit": MAX_PREVIEW_ROWS,
        "_full_rows": data_rows,
    }


# ---------------------------------------------------------------- 各类文件解析
def parse_xlsx(data: bytes, filename: str) -> dict:
    """xlsx/xls → 每 sheet 存 {headers(自动探测切片后), _full_rows(数据区全量),
    _raw_grid(原始全量网格, 供表头行调整后重切), _header_row(自动探测的列名物理行号, 1基)}。"""
    try:
        xls = pd.ExcelFile(io.BytesIO(data))
    except Exception as e:
        raise BizError(f"Excel 文件解析失败: {e}")

    sheets = []
    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name, header=None)
        grid = _grid_rows(df) if not df.empty else []
        header_idx = _detect_header_row(grid) if grid else 0
        headers = grid[header_idx] if grid else []
        data_rows = grid[header_idx + 1:] if grid else []
        m = _sheet_payload(headers, data_rows)
        m["sheet_name"] = sheet_name
        m["_raw_grid"] = grid            # 原始全量网格（含标题行），供 header_row 重切
        m["_header_row"] = header_idx + 1  # 自动探测的列名行物理行号（1 基）
        sheets.append(m)
    if not sheets:
        raise BizError("Excel 文件中没有可用的工作表")
    return {"file_kind": "excel", "sheets": sheets, "sheet_count": len(sheets)}


def _extract_text_lines(text: str) -> list:
    """把解析出的文本按行切分并去除空行。"""
    lines = [ln.strip() for ln in text.splitlines()]
    return [ln for ln in lines if ln]


def _document_sheet(lines: list, auto: dict = None, sheet_name: str = "文档内容") -> dict:
    """把 Word/PDF 文本规范化为与 Excel 相同的单工作表结构。"""
    if auto:
        headers = auto.get("headers") or []
        data_rows = auto.get("rows") or []
    else:
        split_rows = [[c.strip() for c in re.split(r"\t|\|", line)] for line in lines]
        if split_rows and max((len(row) for row in split_rows), default=0) > 1:
            headers, data_rows = split_rows[0], split_rows[1:]
        else:
            headers = ["文本"]
            data_rows = [[line] for line in lines]
    payload = _sheet_payload(headers, data_rows)
    payload.update({
        "sheet_name": sheet_name,
        "_raw_grid": [headers, *data_rows],
        "_header_row": 1,
    })
    return payload


def _regex_prescreen(lines: list, patterns: list) -> list:
    """正则预筛：识别 '姓名：张三' 等模式行，返回二维行数组。"""
    matched = []
    for ln in lines:
        row = []
        ok = True
        for pat in patterns:
            m = re.search(pat, ln)
            if not m:
                ok = False
                break
            row.append(m.group(1).strip())
        if ok:
            matched.append(row)
    return matched


def parse_docx(data: bytes, filename: str) -> dict:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        # 表格行也纳入
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells))
    except Exception as e:
        raise BizError(f"Word 文件解析失败: {e}")

    text = "\n".join(lines)
    cleaned = _extract_text_lines(text)

    # 正则预筛常见模式（姓名：xxx / 学号：xxx / 列:值）
    common_patterns = [
        ([r"姓名[:：]\s*(.+)", r"学号[:：]\s*([0-9]+)"], ["姓名", "学号"]),
        ([r"学号[:：]\s*([0-9]+)", r"姓名[:：]\s*(.+)"], ["学号", "姓名"]),
    ]
    auto = None
    for pats, headers in common_patterns:
        matched = _regex_prescreen(cleaned, pats)
        if len(matched) >= 1:
            auto = {"headers": headers, "rows": matched,
                    "total_rows": len(matched), "preview_limit": MAX_PREVIEW_ROWS}
            break

    sheet = _document_sheet(cleaned, auto, "Word 文档")
    return {
        "file_kind": "docx",
        "text": text,                      # 完整原文（供坐标打标/人工查看）
        "lines": cleaned,
        "auto_parsed": auto,               # 正则预筛结果（可能为 None → 前端坐标打标）
        "sheets": [sheet],
        "sheet_count": 1,
    }


def parse_pdf(data: bytes, filename: str) -> dict:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        lines = []
        for page in reader.pages:
            lines.append((page.extract_text() or ""))
    except Exception as e:
        raise BizError(f"PDF 文件解析失败: {e}")

    text = "\n".join(lines)
    cleaned = _extract_text_lines(text)
    sheet = _document_sheet(cleaned, None, "PDF 文档")
    return {
        "file_kind": "pdf",
        "text": text,
        "lines": cleaned,
        "auto_parsed": None,
        "sheets": [sheet],
        "sheet_count": 1,
    }


def parse_file(data: bytes, filename: str, file_id: str = None) -> dict:
    """统一解析入口：按扩展名分派，返回结果并缓存。"""
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if suffix in ("xlsx", "xls"):
        result = parse_xlsx(data, filename)
    elif suffix == "docx":
        result = parse_docx(data, filename)
    elif suffix == "pdf":
        result = parse_pdf(data, filename)
    else:
        raise BizError(f"不支持的文件类型: .{suffix}（支持 .xlsx/.xls/.docx/.pdf）")

    fid = file_id or make_file_id()
    result["file_id"] = fid
    result["filename"] = filename
    result["md5"] = file_md5(data)
    cache_parsed(fid, result)
    return result
